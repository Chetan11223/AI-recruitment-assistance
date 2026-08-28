import os
from pathlib import Path
from typing import List, Dict, Any, Tuple
import pymupdf as fitz
import docx

class DocumentPage:
    def __init__(self, page_number: int, text: str, char_start_offset: int):
        self.page_number = page_number
        self.text = text
        self.char_start_offset = char_start_offset
        self.char_end_offset = char_start_offset + len(text)

class ParsedDocument:
    def __init__(self, file_name: str, file_type: str, full_text: str, pages: List[DocumentPage]):
        self.file_name = file_name
        self.file_type = file_type
        self.full_text = full_text
        self.pages = pages

    def get_citation_span(self, match_text: str, start_hint: int = -1) -> Dict[str, Any]:
        """Locates character range and page number for an exact or approximate substring."""
        if not match_text or not self.full_text:
            return {"page_number": 1, "char_start": 0, "char_end": 0, "raw_text": match_text}

        cleaned_match = match_text.strip()
        pos = -1
        if start_hint >= 0:
            pos = self.full_text.find(cleaned_match, start_hint)

        if pos == -1:
            pos = self.full_text.find(cleaned_match)

        # If exact match fails, try first 40 chars
        if pos == -1 and len(cleaned_match) > 40:
            pos = self.full_text.find(cleaned_match[:40])

        if pos == -1:
            return {"page_number": 1, "char_start": 0, "char_end": len(match_text), "raw_text": match_text}

        end_pos = pos + len(cleaned_match)
        page_num = 1
        for page in self.pages:
            if page.char_start_offset <= pos <= page.char_end_offset:
                page_num = page.page_number
                break

        return {
            "page_number": page_num,
            "char_start": pos,
            "char_end": end_pos,
            "raw_text": cleaned_match
        }

class DocumentLoader:
    @staticmethod
    def load_pdf(file_path: str) -> ParsedDocument:
        doc = fitz.open(file_path)
        pages: List[DocumentPage] = []
        full_text_parts = []
        current_offset = 0

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            page_text = page.get_text("text") or ""
            # Normalize whitespace while preserving linebreaks
            pages.append(DocumentPage(
                page_number=page_idx + 1,
                text=page_text,
                char_start_offset=current_offset
            ))
            full_text_parts.append(page_text)
            current_offset += len(page_text) + 1  # +1 for newline between pages

        doc.close()
        full_text = "\n".join(full_text_parts)
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="pdf",
            full_text=full_text,
            pages=pages
        )

    @staticmethod
    def load_docx(file_path: str) -> ParsedDocument:
        doc = docx.Document(file_path)
        paragraphs_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs_text.append(p.text)

        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs_text.append(row_text)

        full_text = "\n".join(paragraphs_text)
        pages = [DocumentPage(page_number=1, text=full_text, char_start_offset=0)]
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="docx",
            full_text=full_text,
            pages=pages
        )

    @staticmethod
    def load_text(file_path: str) -> ParsedDocument:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()
        pages = [DocumentPage(page_number=1, text=full_text, char_start_offset=0)]
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="txt",
            full_text=full_text,
            pages=pages
        )

    @classmethod
    def load_file(cls, file_path: str) -> ParsedDocument:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return cls.load_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return cls.load_docx(file_path)
        else:
            return cls.load_text(file_path)
